"""
Plain-text extraction for uploaded Documents, so the AI assistant can read
their content on demand (via a tool call) instead of the app needing to
guess ahead of time which document might be relevant to a question.

Supported: .xlsx/.xls, .csv, .txt, .docx, .pdf
Unsupported types (images, ...) are skipped gracefully — the file is
still stored and downloadable, it just isn't readable by the assistant.
"""
import csv

import openpyxl

_MAX_CHARS = 20_000  # keep a single document from blowing out the LLM context


def _truncate(text: str) -> str:
    if len(text) <= _MAX_CHARS:
        return text
    return text[:_MAX_CHARS] + f'\n\n[... truncated, {len(text) - _MAX_CHARS} more characters not shown ...]'


def _extract_xlsx(filepath: str) -> str:
    wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f'## Sheet: {ws.title}')
        for row in ws.iter_rows(values_only=True):
            if row is None or all(v is None for v in row):
                continue
            cells = ['' if v is None else str(v) for v in row]
            parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _extract_csv(filepath: str) -> str:
    with open(filepath, newline='', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        return '\n'.join(' | '.join(row) for row in reader)


def _extract_txt(filepath: str) -> str:
    with open(filepath, encoding='utf-8', errors='replace') as f:
        return f.read()


def _extract_docx(filepath: str) -> str:
    import docx
    doc = docx.Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text for cell in row.cells))
    return '\n'.join(parts)


def _extract_pdf(filepath: str) -> str:
    import pypdf
    reader = pypdf.PdfReader(filepath)
    parts = [page.extract_text() or '' for page in reader.pages]
    return '\n'.join(parts)


_EXTRACTORS = {
    'xlsx': _extract_xlsx,
    'xlsm': _extract_xlsx,
    'xls': _extract_xlsx,
    'csv': _extract_csv,
    'txt': _extract_txt,
    'docx': _extract_docx,
    'pdf': _extract_pdf,
}


def extract_text(filepath: str, file_type: str) -> tuple[str | None, str | None]:
    """Returns (extracted_text, error). Exactly one will be None.

    error is a short, user-facing reason extraction didn't happen — either
    an unsupported file type, or a genuine parse failure (corrupt/unreadable
    file), never a raised exception.
    """
    ext = (file_type or '').lower().lstrip('.')
    extractor = _EXTRACTORS.get(ext)
    if not extractor:
        return None, f'Text extraction is not supported for .{ext} files.'
    try:
        text = extractor(filepath)
        if not text.strip():
            return None, 'No readable text content was found in this file.'
        return _truncate(text), None
    except Exception as exc:
        return None, f'Could not extract text: {exc}'
